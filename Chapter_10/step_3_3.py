from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Mm

from step_1_1 import IMG_DIR, OUT_DIR  
from step_2_1 import OUT_2_1
from step_3_1 import apply_font
from step_3_2 import OUT_3_2, add_blank_paragraph

OUT_3_3 = OUT_DIR / f"{Path(__file__).stem}.docx"


def insert_indicators():
    doc = Document(OUT_3_2)  
    r_head = doc.add_paragraph().add_run("1. 주요 금리 현황")  
    apply_font(r_head, size_pt=14, is_bold=True)  
    add_blank_paragraph(doc, size_pt=10) 

    table = doc.add_table(rows=1, cols=5) 
    table.alignment = WD_TABLE_ALIGNMENT.CENTER 
    table.allow_autofit = False 

    tr = table.rows[0]  
    with pd.ExcelFile(OUT_2_1) as xlsx:  
        for idx, sheet_name in enumerate(xlsx.sheet_names):  
            df_raw = pd.read_excel(xlsx, sheet_name=sheet_name) 
            df_raw = df_raw.tail(24)  
            td = tr.cells[idx] 
            td.width = Mm(35.5)  

            p1 = td.paragraphs[0] 
            r1 = p1.add_run(sheet_name) 
            apply_font(r1, size_pt=12, is_bold=True, rgb="333333")

            p2 = td.add_paragraph()  
            last_value = df_raw["DATA_VALUE"].iloc[-1]  
            r2 = p2.add_run(f"{last_value:,.2f}") 
            apply_font(r2, size_pt=14, is_bold=True, rgb="333333")

            p3 = td.add_paragraph()  
            diff = last_value - df_raw["DATA_VALUE"].iloc[0]  
            arrow = "▲" if diff > 0 else "▼" if diff < 0 else ""  
            rgb = "FF0000" if diff > 0 else "0000FF" if diff < 0 else "000000"  
            r3 = p3.add_run(f"{arrow}{abs(diff):,.2f}%p")  
            apply_font(r3, size_pt=10, is_bold=True, rgb=rgb)

            p4 = td.add_paragraph()  
            p4.paragraph_format.left_indent = Mm(-1)  
            img_path = IMG_DIR / f"{sheet_name}.png"  
            p4.add_run().add_picture(img_path.as_posix(), Mm(30), Mm(8)) 

    add_blank_paragraph(doc, size_pt=10)  
    doc.save(OUT_3_3)  


if __name__ == "__main__":
    insert_indicators()  